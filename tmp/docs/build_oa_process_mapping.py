import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path('/Users/mac/项目/若伊部署')
CSV_PATH = Path('/Users/mac/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_tc376xz27m2222_cb11/temp/RWTemp/2026-06/8356ac11f3ee5f741aa545a5179604e7/20260608152558TxXJ.csv')
OUT_DIR = BASE / 'output' / 'doc'
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / '规划院OA用户与审批流程映射表-20260710.docx'


APPROVALERS = {
    '侯斌超': '侯斌超（院长）',
    '钱爱梅': '钱爱梅（分管领导）',
    '邱燕萍': '邱燕萍（分管领导）',
    '马倩': '马倩（分管领导/所长）',
    '张华': '张华（总工/分管领导）',
    '张龄': '张龄（总师）',
    '朱新捷': '朱新捷（总师）',
    '濮卫民': '濮卫民（城市设计所所长）',
    '邱燕萍所长': '邱燕萍（总体规划和专项规划所所长）',
    '李丹': '李丹（乡村规划所所长）',
    '罗翔': '罗翔（规划研究室部门领导）',
    '缪云涛': '缪云涛（交通市政规划所部门领导）',
    '苏莉': '苏莉（综合科人事/行政/财务负责人）',
    '何志华': '何志华（综合科经营部门领导）',
    '严己': '严己（技术审查室节点）',
    '赖志勇': '赖志勇（技术审查室特殊节点）',
}


def route(*names):
    return ' → '.join(APPROVALERS.get(n, n) for n in names)


ROUTES = {
    '城市更新设计所': {
        'unit': '城市更新设计所（架构图称“城市更新规划所”）',
        'type': '普通所流程',
        'approvers': route('马倩', '钱爱梅', '侯斌超'),
        'note': 'Word问题记录中的毛丹案例明确为：马倩 → 钱爱梅 → 侯斌超。',
    },
    '城市设计所': {
        'unit': '城市设计所',
        'type': '普通所流程',
        'approvers': route('濮卫民', '钱爱梅', '侯斌超'),
        'note': '按组织架构图中“濮卫民 @城市设计所”映射。',
    },
    '总体规划和专项规划所': {
        'unit': '总体规划和专项规划所',
        'type': '普通所流程',
        'approvers': route('邱燕萍所长', '钱爱梅', '侯斌超'),
        'note': '按组织架构图中“邱燕萍 @总体规划和专项规划所”映射。',
    },
    '乡村规划所': {
        'unit': '乡村规划所',
        'type': '普通所流程',
        'approvers': route('李丹', '钱爱梅', '侯斌超'),
        'note': '按组织架构图中“李丹 @乡村规划所”映射。',
    },
    '规划研究室': {
        'unit': '规划研究室',
        'type': '部门流程',
        'approvers': route('罗翔', '马倩', '侯斌超'),
        'note': '赖志勇属于该部门，但架构图另标有其专项路径，详见特殊流程。',
    },
    '数字城市规划所': {
        'unit': '数字城市规划所',
        'type': '部门流程（负责人待补）',
        'approvers': '部门负责人未配置（架构图为空） → ' + route('马倩', '侯斌超'),
        'note': '必须先核实数字城市规划所部门负责人是否缺失；当前不能按自动通过处理。',
    },
    '交通市政规划所': {
        'unit': '交通市政规划所',
        'type': '总工流程',
        'approvers': route('缪云涛', '张华', '侯斌超'),
        'note': '按组织架构图中“缪云涛 @交通市政规划所”进入张华节点。',
    },
    '综合科（人事/行政/财务）': {
        'unit': '综合科（人事/行政/财务）',
        'type': '综合科流程',
        'approvers': route('苏莉', '邱燕萍', '侯斌超'),
        'note': '图中表格明确列出综合科（人事、行政、财务）负责人为苏莉。',
    },
    '综合科（经营）': {
        'unit': '综合科（经营）',
        'type': '经营流程',
        'approvers': route('何志华', '张华', '侯斌超'),
        'note': '架构图另标李豪杰专项路径：李豪杰 → 何志华 → 张华 → 侯斌超。',
    },
    '技术审查室': {
        'unit': '技术审查室',
        'type': '特殊流程',
        'approvers': '严己节点 → 张华、张龄、朱新捷三选一 → 侯斌超',
        'note': '架构图标注“三选一”；不能把该节点自动通过。赖志勇另走朱新捷。',
    },
}


SHARED_DEPTS = {'公共空间', '科研小组', '专家库', '案例库', '院部'}


def parse_rows():
    with CSV_PATH.open(encoding='utf-8-sig', newline='') as f:
        rows = list(csv.reader(f))
    result = []
    for row in rows[1:]:
        idx, account, role, storage, dept_text = row
        name = account.split('/')[0]
        depts = [x.strip() for x in dept_text.split(';') if x.strip()]
        result.append({'id': idx, 'account': account, 'name': name, 'role': role, 'storage': storage, 'depts': depts})
    return result


def display_depts(depts):
    business = [d for d in depts if d not in SHARED_DEPTS]
    if not business:
        return '无规划院业务所/部门（仅公共空间或测试部门）'
    normalized = []
    for d in business:
        if d in {'综合科', '行政办公', '计划经营', '财务管理'}:
            label = '综合科（含人事/行政/财务/经营）'
        elif d == '城市更新设计所':
            label = '城市更新设计所（架构图称“城市更新规划所”）'
        else:
            label = d
        if label not in normalized:
            normalized.append(label)
    return '、'.join(normalized)


def multi_route(depts):
    routes = []
    if any(d in depts for d in ['城市更新设计所', '城市设计所', '总体规划和专项规划所', '乡村规划所']):
        labels = []
        if '城市更新设计所' in depts: labels.append('城市更新：马倩 → 钱爱梅 → 侯斌超')
        if '城市设计所' in depts: labels.append('城市设计：濮卫民 → 钱爱梅 → 侯斌超')
        if '总体规划和专项规划所' in depts: labels.append('总体规划：邱燕萍 → 钱爱梅 → 侯斌超')
        if '乡村规划所' in depts: labels.append('乡村规划：李丹 → 钱爱梅 → 侯斌超')
        routes.append('；'.join(labels))
    if '规划研究室' in depts: routes.append('规划研究室：罗翔 → 马倩 → 侯斌超')
    if '数字城市规划所' in depts: routes.append('数字城市：部门负责人待补 → 马倩 → 侯斌超')
    if '交通市政规划所' in depts: routes.append('交通市政：缪云涛 → 张华 → 侯斌超')
    if any(d in depts for d in ['综合科', '行政办公', '计划经营', '财务管理']):
        routes.append('综合科：按人事/行政/财务或经营分别走苏莉/何志华 → 对应分管节点 → 侯斌超')
    if '技术审查室' in depts: routes.append('技术审查：张华/张龄/朱新捷三选一 → 侯斌超')
    return '；'.join(routes)


def classify(rec):
    name = rec['name']
    depts = rec['depts']
    role = rec['role']
    if name == '侯斌超':
        return '院长/审批终点账号', '院部（院长）', '审批终点账号，不作为普通发起人；如需用该账号发起，请单独确认上级节点。', '保留审批账号'
    if name == '马倩':
        return '所长/分管领导账号', display_depts(depts), ROUTES['城市更新设计所']['approvers'], '保留审批账号，同时验证其作为所长和分管领导的节点身份。'
    if name == '朱新捷':
        return '总师/技术审查特殊账号', '技术审查室', '如由该账号发起，按架构图“直接到院长”：侯斌超（院长）', '保留审批账号'
    if name == '严己':
        return '技术审查特殊账号', '技术审查室', '张华、张龄、朱新捷三选一 → 侯斌超（院长）', '需验证三选一是否显示具体人名'
    if name == '赖志勇':
        return '技术审查特殊账号', '规划研究室（兼专家库）', '朱新捷（总师） → 侯斌超（院长）', '架构图明确标注“赖志勇到朱新捷”'
    if name in {'金晨', '黄丽萍'}:
        return '多部门/特殊账号', display_depts(depts), multi_route(depts), '同一账号必须按实际发起部门分别走流程；不能只验证一条路线。'
    if name in {'系统管理员', '751i扫描用户', 'demo', '测试用户2'} or not any(d not in SHARED_DEPTS for d in depts):
        return '非OA业务账号', display_depts(depts), '不纳入规划院OA审批主流程；如需验证登录或权限，单独做账号可用性测试。', '不作为流程发起人'
    if '技术审查室' in depts:
        return '技术审查普通/特殊流程', display_depts(depts), ROUTES['技术审查室']['approvers'], '技术审查节点按图示三选一，需人工登录审批。'
    preferred = ['城市更新设计所', '城市设计所', '总体规划和专项规划所', '乡村规划所', '规划研究室', '数字城市规划所', '交通市政规划所']
    selected = next((d for d in preferred if d in depts), None)
    if selected:
        r = ROUTES[selected]
        return r['type'], display_depts(depts), r['approvers'], r['note']
    if any(d in depts for d in ['综合科', '行政办公', '计划经营', '财务管理']):
        key = '综合科（经营）' if '计划经营' in depts else '综合科（人事/行政/财务）'
        r = ROUTES[key]
        return r['type'], display_depts(depts), r['approvers'], r['note']
    return '待核实', display_depts(depts), 'CSV中有业务部门，但无法从架构图唯一确定审批链。', '证据不足，需补充部门负责人配置'


DOC_FONT = 'PingFang SC'


def set_run_font(run, size=8.5, bold=False, color='1F2937'):
    run.font.name = DOC_FONT
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), DOC_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), DOC_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), DOC_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.insert(0, tblW)
    tblW.set(qn('w:w'), str(sum(widths_dxa)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def style_table(table, header_fill='DCE6F1', body_size=7.4):
    set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_shading(cell, header_fill if ri == 0 else 'FFFFFF')
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                if ri == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=7.8 if ri == 0 else body_size, bold=(ri == 0), color='17365D' if ri == 0 else '1F2937')


def add_table(doc, headers, rows, widths_dxa, body_size=7.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = str(text)
    set_table_geometry(table, widths_dxa)
    style_table(table, body_size=body_size)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size={1: 14, 2: 11.5, 3: 10}[level], bold=True, color='2E74B5' if level < 3 else '1F4D78')
    return p


def build_doc(records):
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.1)
    sec.bottom_margin = Cm(1.1)
    sec.left_margin = Cm(1.1)
    sec.right_margin = Cm(1.1)
    sec.header_distance = Cm(0.6)
    sec.footer_distance = Cm(0.6)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), DOC_FONT)
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1
    for sname, size, color, before, after in [
        ('Heading 1', 14, '2E74B5', 10, 5),
        ('Heading 2', 11.5, '2E74B5', 7, 4),
        ('Heading 3', 10, '1F4D78', 5, 3),
    ]:
        st = styles[sname]
        st.font.name = DOC_FONT
        st._element.rPr.rFonts.set(qn('w:eastAsia'), DOC_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('规划院 OA 流程测试用户与审批流程映射表')
    set_run_font(r, size=18, bold=True, color='17365D')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('编制日期：2026年7月10日｜范围：可道云 CSV 中 50 个账号')
    set_run_font(r, size=8.5, color='6B7280')

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(5)
    run = note.add_run('使用口径：')
    set_run_font(run, size=8.5, bold=True, color='7A5A00')
    run = note.add_run('表中“审批经过人员”只列审批节点，不计发起人本人。流程按组织架构图和 Word 问题记录整理；每个节点必须使用对应可道云账号登录后人工审批，不能按“系统内已有用户”直接自动通过。')
    set_run_font(run, size=8.5, color='7A5A00')

    add_heading(doc, '一、组织架构对应的流程规则', 1)
    route_rows = []
    for key, r in ROUTES.items():
        route_rows.append([r['unit'], r['type'], r['approvers'], r['note']])
    add_table(doc, ['所/部门', '流程类型', '审批经过人员', '证据与测试提示'], route_rows, [2100, 1250, 3800, 2210], body_size=7.5)

    add_heading(doc, '二、CSV账号逐用户映射', 1)
    user_rows = []
    for rec in records:
        flow_type, unit, approvers, note = classify(rec)
        user_rows.append([rec['id'], rec['name'], rec['role'], unit, flow_type, approvers, note])
    add_table(doc, ['编号', '用户', '可道云角色', '所属所/部门', '测试分类', '审批经过人员', '备注/测试要求'], user_rows, [480, 1050, 1050, 2400, 1250, 3350, 1950], body_size=6.6)

    add_heading(doc, '三、架构图中需补充确认或单独保留的账号', 1)
    extra_rows = [
        ['张华', '总工/分管领导', '架构图明确存在；CSV未找到同名账号', '部分特殊流程节点；图中另标“直接到院长”', '确认可道云账号存在并可登录审批'],
        ['张龄', '总师', '架构图明确存在；CSV未找到同名账号', '技术审查室“三选一”节点；图中另标“直接到院长”', '确认可道云账号存在并可登录审批'],
        ['李豪杰', '综合科经营特殊节点', '架构图明确存在；CSV未找到同名账号', '李豪杰 → 何志华 → 张华 → 侯斌超', '需要补建或确认账号后专项验证'],
        ['数字城市规划所负责人', '部门负责人', '架构图对应位置为空', '负责人节点 → 马倩 → 侯斌超', '补齐负责人前，不应把节点视为自动通过'],
    ]
    add_table(doc, ['账号/节点', '角色', '材料状态', '审批路径', '动作'], extra_rows, [1450, 1500, 2800, 3000, 2200], body_size=7.4)

    add_heading(doc, '四、建议测试执行顺序', 1)
    steps = [
        '先用普通用户分别覆盖城市更新设计所、城市设计所、总体规划和专项规划所、乡村规划所、规划研究室、交通市政规划所，确认标准链路能完整到院长。',
        '再覆盖数字城市规划所、综合科人事/行政/财务、综合科经营、技术审查室等负责人为空或存在特殊分支的部门。',
        '对每个发起人：可道云中已有审批账号时，登录对应账号完成审批；没有审批账号时，先建立对应账号再继续，不能自动通过。',
        '一条流程完整结束且无报错后，可删除本次普通发起人账号，再创建下一名对应用户继续测试；院长、分管领导、所长、总工/总师等审批节点账号建议保留。',
        '每次测试记录：发起人、所属所/部门、实际经过的审批人、表单内容是否可见、节点是否显示具体人名、是否出现自动通过或节点错位。',
    ]
    for idx, text in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f'{idx}. {text}')
        set_run_font(r, size=8.8, color='1F2937')

    add_heading(doc, '五、来源与证据边界', 1)
    sources = [
        '规划院OA流程测试问题-20260710-1.docx：包含毛丹案例、人工审批要求和问题描述。',
        '20260608152558TxXJ.csv：可道云账号、角色和所在部门，共50条账号记录。',
        '653e8c5eeb8ef0a5fa6e301d22c53291.png、4eff97b33ba725d0c0d8a3b20507970a.png：规划院组织架构及审批流程图。',
        '“数字城市规划所负责人为空”、CSV中没有张华/张龄/李豪杰等事项，属于材料中可见的待补齐项，不在本表中虚构账号状态。',
    ]
    for text in sources:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run('• ' + text)
        set_run_font(r, size=8.2, color='4B5563')

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run('规划院 OA 流程测试映射表')
    set_run_font(r, size=7.5, color='9CA3AF')

    doc.save(OUT_PATH)
    print(OUT_PATH)
    print('records=', len(records))


if __name__ == '__main__':
    build_doc(parse_rows())
